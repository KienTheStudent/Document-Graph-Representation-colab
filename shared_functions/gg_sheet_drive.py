'''
To use these functions, grant the "Editor" permission to the service account email on the target Google Sheet/ Google Drive and
                                                                            change the ID in the function parameter

"gg-service-agent@legalrag-471601.iam.gserviceaccount.com" 

Link to sheet: https://docs.google.com/spreadsheets/d/1xBgBiA1KwTNdqPfrqH5p_Sf-MhTCXMfy4ousb0WE4Ik/edit

Link to drive: https://drive.google.com/drive/folders/16dRxPz4tVDPScwuYOQ_U5opRdg96qW9W?usp=drive_link
'''
import io, sys, os
import gspread
import pandas as pd
import json
# from pydrive.auth import GoogleAuth
# from pydrive.drive import GoogleDrive
# from oauth2client.service_account import ServiceAccountCredentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

from typing import List, Dict, Optional
from google.oauth2 import service_account
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from collections import defaultdict
from dotenv import load_dotenv
import mimetypes

load_dotenv()

google_api_creds = '/content/ggsheet_credentials.json'
spreadsheet_id = os.getenv('GOOGLE_SHEET_ID') 
drive_id = os.getenv('GOOGLE_DRIVE_ID')

#Google Sheet
def gs_to_df_pandas(tab_name, spreadsheet_id = spreadsheet_id, creds_path=google_api_creds):
  gc = gspread.service_account(filename=creds_path)
  sh = gc.open_by_key(spreadsheet_id)
  wks = sh.worksheet(tab_name)
  df = pd.DataFrame(wks.get_all_records())
  return df

def gs_to_dict( tab_name, spreadsheet_id = spreadsheet_id, creds_path=google_api_creds):
    gc = gspread.service_account(filename=creds_path)
    sh = gc.open_by_key(spreadsheet_id)
    wks = sh.worksheet(tab_name)
    results_json = wks.get_all_records()
    return results_json

def write_df_to_gs(df, tab_name, spreadsheet_id = spreadsheet_id, creds_path=google_api_creds):
    import gspread
    from gspread.exceptions import WorksheetNotFound
    
    gc = gspread.service_account(filename=creds_path)
    sh = gc.open_by_key(spreadsheet_id)

    try:
        # Try to open existing worksheet
        wks = sh.worksheet(tab_name)
        # Find the last filled row
        last_row = len(wks.get_all_values())
        # Append DataFrame values (without header)
        wks.update(f"A{last_row+1}", df.values.tolist())
        return f"DataFrame appended to existing Google Sheet tab: {tab_name}"
    except WorksheetNotFound:
        # If worksheet does not exist, create it
        wks = sh.add_worksheet(title=tab_name, rows=str(len(df)+1), cols=str(len(df.columns)))
        wks.update([df.columns.values.tolist()] + df.values.tolist())
        return f"New tab created and DataFrame written to Google Sheet: {tab_name}"

#Google Drive
def get_drive_service(creds_path: str = google_api_creds, scopes: Optional[List[str]] = None):
    """
    Authenticate and return a Google Drive API service object.
    """
    if scopes is None:
        scopes = ['https://www.googleapis.com/auth/drive']
    creds = service_account.Credentials.from_service_account_file(creds_path, scopes=scopes)
    return build('drive', 'v3', credentials=creds)

def list_drive_files(
    folder_id: str = drive_id,
    creds_path: str = google_api_creds,
    prefix_path: str = ''):
    
    service = get_drive_service(creds_path)
    results = []

    def safe_list_files(**kwargs):
        """Handle 500 errors with retries and pagination."""
        retries = 5
        delay = 1
        while retries > 0:
            try:
                return service.files().list(**kwargs).execute()
            except HttpError as e:
                if e.resp.status == 500:
                    print(f"500 error, retrying in {delay}s...")
                    time.sleep(delay)
                    delay *= 2
                    retries -= 1
                else:
                    raise
        raise RuntimeError("Failed after multiple retries due to 500 errors.")

    def _list_recursive(current_id: str, current_path: str):
        page_token = None
        while True:
            query = f"'{current_id}' in parents and trashed = false"
            response = safe_list_files(
                q=query,
                fields="nextPageToken, files(id, name, mimeType)",
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
                pageToken=page_token,
                pageSize=100
            )
            items = response.get('files', [])
            for item in items:
                full_path = os.path.join(current_path, item['name'])
                results.append({
                    'id': item['id'],
                    'name': item['name'],
                    'path': full_path,
                    'mimeType': item['mimeType']
                })
                if item['mimeType'] == 'application/vnd.google-apps.folder':
                    _list_recursive(item['id'], full_path)

            page_token = response.get('nextPageToken')
            if not page_token:
                break

    _list_recursive(folder_id, prefix_path)

    # --- Build nested tree ---
    def build_tree(paths):
        tree = lambda: defaultdict(tree)
        root = tree()
        for p in paths:
            parts = p.split(os.sep)
            node = root
            for part in parts:
                node = node[part]
        return root

    def print_tree(node, prefix=""):
        items = sorted(node.keys())
        for i, key in enumerate(items):
            connector = "└── " if i == len(items) - 1 else "├── "
            print(prefix + connector + key)
            print_tree(node[key], prefix + ("    " if i == len(items) - 1 else "│   "))

    def show_drive_tree(file_list):
        paths = [f["path"] for f in file_list]
        tree = build_tree(paths)
        print_tree(tree)  # return the flat list if needed

    return show_drive_tree(results)

def find_file_full_path(
    filename: str,
    creds_path: str = google_api_creds,
    drive_id: Optional[str] = None
) -> Optional[Dict]:

    creds = service_account.Credentials.from_service_account_file(
        creds_path,
        scopes=["https://www.googleapis.com/auth/drive"]
    )
    service = build("drive", "v3", credentials=creds)

    # === 2. Search for file ===
    q = f"name='{filename}' and trashed=false"
    params = {
        "q": q,
        "fields": "files(id, name, parents)",
        "supportsAllDrives": True,
        "includeItemsFromAllDrives": True,
    }
    if drive_id:
        params.update({"corpora": "drive", "driveId": drive_id})

    response = service.files().list(**params).execute()
    files = response.get("files", [])

    if not files:
        print(f"❌ File '{filename}' not found.")
        return None

    file = files[0]
    path_parts = [file["name"]]
    parent_id = file.get("parents", [None])[0]

    # === 3. Walk upward to reconstruct full folder path ===
    while parent_id:
        parent = service.files().get(
            fileId=parent_id,
            fields="id, name, parents",
            supportsAllDrives=True
        ).execute()
        path_parts.insert(0, parent["name"])
        parent_id = parent.get("parents", [None])[0]

    full_path = "/".join(path_parts[2:])

    # === 4. Return consistent dictionary ===
    return full_path

def read_drive_file(path: str, creds_path: str = google_api_creds, as_type: str = None, drive_id: str = drive_id):
    
    full_path = path if "/" in path else find_file_full_path(path)
    service = get_drive_service(creds_path)
    parts = full_path.replace("\\", "/").split("/")
    parent_id = drive_id

    # Navigate folder structure
    for i, part in enumerate(parts):
        q = f"'{parent_id}' in parents and name = '{part}' and trashed = false"
        results = service.files().list(q=q, fields="files(id, name, mimeType)").execute()
        files = results.get("files", [])
        if not files:
            raise FileNotFoundError(f"❌ Path component not found: {part} (full: {'/'.join(parts[:i+1])})")
        file = files[0]
        parent_id = file["id"]

    if file["mimeType"] == "application/vnd.google-apps.folder":
        raise IsADirectoryError(f"'{full_path}' is a folder, not a file")

    file_name = file["name"]
    mime_type = file["mimeType"]

    # Infer file type if not specified
    if as_type is None:
        if file_name.endswith(".json") or "json" in mime_type:
            as_type = "json"
        elif file_name.endswith(".csv") or "csv" in mime_type:
            as_type = "csv"
        elif file_name.endswith(".docx") or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            as_type = "docx"
        elif file_name.endswith(".doc") or mime_type == "application/msword":
            as_type = "doc"
        elif file_name.endswith(".pdf") or "pdf" in mime_type:
            as_type = "pdf"
        else:
            as_type = "txt"

    # Download file (raw content for all except Google Docs files)
    request = service.files().get_media(fileId=file["id"])
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    fh.seek(0)

    # Parse based on inferred type
    if as_type == "json":
        return json.load(io.TextIOWrapper(fh, encoding="utf-8"))

    elif as_type == "csv":
        return pd.read_csv(fh)

    elif as_type == "txt":
        try:
            return fh.read().decode("utf-8")
        except UnicodeDecodeError:
            raise UnicodeDecodeError("❌ Unable to decode as UTF-8. Try using as_type='doc', 'docx', or 'pdf'.")

    elif as_type == "docx":
        try:
            from docx import Document
            doc = Document(fh)
            full_text = []

            # Extract all paragraph text
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            return "\n".join([t for t in full_text if t]).strip()
        except ImportError:
            raise ImportError("📦 Missing dependency: install python-docx (`pip install python-docx`)")


    elif as_type == "doc":
        try:
            import mammoth  # pip install mammoth
            result = mammoth.extract_raw_text(fh)
            return result.value.strip()
        except ImportError:
            raise ImportError("📦 Missing dependency: install mammoth (`pip install mammoth`)")

    elif as_type == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(fh)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text.strip()
        except ImportError:
            raise ImportError("📦 Missing dependency: install PyPDF2 (`pip install PyPDF2`)")

    else:
        raise ValueError(f"Unsupported file type: {as_type}")

def find_folder_by_name(
    folder_name: str,
    creds_path: str = google_api_creds,
    drive_id: Optional[str] = None
) -> Optional[Dict]:
    """
    Find a folder by name and return:
    {
        'id': folder_id,
        'name': folder_name,
        'full_path': 'A/B/C'
    }
    """
    creds = service_account.Credentials.from_service_account_file(
        creds_path,
        scopes=["https://www.googleapis.com/auth/drive"]
    )
    service = build("drive", "v3", credentials=creds)

    q = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    
    params = {
        "q": q,
        "fields": "files(id, name, parents)",
        "supportsAllDrives": True,
        "includeItemsFromAllDrives": True,
    }
    if drive_id:
        params.update({"corpora": "drive", "driveId": drive_id})

    response = service.files().list(**params).execute()
    folders = response.get("files", [])

    if not folders:
        print(f"❌ Folder '{folder_name}' not found.")
        return None

    folder = folders[0]
    folder_id = folder["id"]

    # Build full path (same logic as your function)
    path_parts = [folder["name"]]
    parent_id = folder.get("parents", [None])[0]

    while parent_id:
        parent = service.files().get(
            fileId=parent_id,
            fields="id, name, parents",
            supportsAllDrives=True
        ).execute()
        path_parts.insert(0, parent["name"])
        parent_id = parent.get("parents", [None])[0]

    full_path = "/".join(path_parts)  # Remove root

    return {
        "id": folder_id,
        "name": folder_name,
        "full_path": full_path
    }

def count_files_in_folder(
    folder_id: str,
    creds_path: str = google_api_creds,
    recursive: bool = False
):
    """
    Count how many files are inside a specific Google Drive folder
    and list their names. Validates folder existence.
    """
    service = get_drive_service(creds_path)

    # --- Validate folder ID first ---
    try:
        meta = service.files().get(
            fileId=folder_id,
            fields="id, name, mimeType",
            supportsAllDrives=True
        ).execute()
    except Exception as e:
        raise ValueError(f"❌ Invalid folder ID: '{folder_id}'. Google Drive returned: {e}")

    if meta["mimeType"] != "application/vnd.google-apps.folder":
        raise ValueError(f"❌ The ID you provided is not a folder: {folder_id}")

    file_names = []

    def _scan(current_id: str):
        query = f"'{current_id}' in parents AND trashed = false"
        response = service.files().list(
            q=query,
            fields="files(id, name, mimeType)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()

        items = response.get("files", [])

        for item in items:
            if item["mimeType"] == "application/vnd.google-apps.folder":
                if recursive:
                    _scan(item["id"])
            else:
                file_names.append(item["name"])

    _scan(folder_id)
    return len(file_names), file_names

def count_files_by_folder_name(
    folder_name: str,
    creds_path: str = google_api_creds,
    recursive: bool = False
):
    """
    Count files inside a folder specified by *name*, not ID.
    """
    folder_info = find_folder_by_name(folder_name, creds_path)

    if folder_info is None:
        return 0, []

    folder_id = folder_info["id"]

    # Reuse your folder-counting logic:
    return count_files_in_folder(folder_id, creds_path, recursive)

def upload_file_to_drive(
    filepath: str,
    folder_name: str,
    creds_path: str = google_api_creds
):
    """
    Upload a local file to a Google Drive folder specified by name.
    Supports Shared Drives for service accounts.

    Parameters:
        filepath (str): Path to the local file to upload.
        folder_name (str): Name of the target Google Drive folder (can be in a Shared Drive).
        creds_path (str): Path to service account credentials JSON.
    """

    # --- 1. Find folder by name ---
    folder_info = find_folder_by_name(folder_name, creds_path)
    if folder_info is None:
        raise ValueError(f"❌ Folder '{folder_name}' not found on Google Drive.")
    
    folder_id = folder_info["id"]

    # --- 2. Prepare file metadata ---
    file_name = os.path.basename(filepath)
    mime_type, _ = mimetypes.guess_type(filepath)
    if mime_type is None:
        mime_type = "application/octet-stream"

    file_metadata = {
        "name": file_name,
        "parents": [folder_id]
    }

    # --- 3. Upload file ---
    service = get_drive_service(creds_path)
    media = MediaFileUpload(filepath, mimetype=mime_type, resumable=True)
    
    try:
        uploaded_file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id, name",
            supportsAllDrives=True  # <-- important for Shared Drives
        ).execute()
    except HttpError as e:
        raise RuntimeError(f"❌ Failed to upload file: {e}")

    full_path = os.path.join(folder_info.get("full_path", folder_name), file_name)

    print(f"✅ File uploaded: {full_path} to Folder {folder_name}")

