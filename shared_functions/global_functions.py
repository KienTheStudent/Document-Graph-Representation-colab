import pandas as pd
import numpy as np
import io
import boto3
from botocore.exceptions import NoCredentialsError
# from docx import Document
import PyPDF2
# import pymysql
from neo4j import GraphDatabase
# from supabase import create_client, Client
from typing import List, Dict, Any, Optional
from docx import Document
import tempfile
from docx2pdf import convert
import subprocess
import os
import win32com.client as win32
import mammoth, pdfplumber, tempfile, docx2txt
import json

s3 = boto3.client("s3")

#For S3 file call
def list_files_recursive(bucket_name = 'legaldocstorage', file_types = None) -> list[str]:
    '''
    List all files inside the S3 bucket in a recursive manner (files and folders)
    '''
    paginator = s3.get_paginator("list_objects_v2")

    files = []
    for page in paginator.paginate(Bucket=bucket_name):
        for obj in page.get("Contents", []):
            files.append(obj["Key"])
    
    if file_types is None:
        return [f for f in files]
    else:
        return [f for f in files if f.lower().endswith(tuple(f".{ft}" for ft in file_types))]

def upload_file_to_s3(file_path: str, bucket_name: str = 'legaldocstorage', expire: int = 3600) -> str:
    """
    Upload a file to S3 and return a presigned URL.
    
    Parameter:
    
    file_path: local path to the target file
    
    """
    
    object_name = file_path.split("/")[-1]
    if object_name is None:
        object_name = file_path.split("/")[-1]  # default: file name

    try:
        s3.upload_file(file_path, bucket_name, object_name)
        print(f"✅ Uploaded {file_path} to s3://{bucket_name}/{object_name}")

        # Generate presigned URL
        url = s3.generate_presigned_url(
            "get_object",
            Params={"Bucket": bucket_name, "Key": object_name},
            ExpiresIn=expire,
        )
        return f'{bucket_name}/{object_name}'
    except FileNotFoundError:
        raise Exception("❌ The file was not found")
    except NoCredentialsError:
        raise Exception("❌ AWS credentials not available")

def bucket_object_separator(bucket_object: str) -> tuple[str, str]:
    """
    Separates the bucket name and object name from a bucket object string.
    """
    bucket_name, object_name = bucket_object.split("/", 1)
    return bucket_name, object_name

def download_file_from_s3(bucket_object: str, local_path: str = ''):
    """
    Downloads whole file from S3 and saves it locally.
    """
    bucket_name, object_name = bucket_object_separator(bucket_object)
    try:
        s3.download_file(bucket_name, object_name, local_path)
        print(f"✅ Downloaded s3://{bucket_name}/{object_name} to {local_path}")
    except NoCredentialsError:
        raise Exception("❌ AWS credentials not available")

def extract_docx_all_text(body):
    '''
    Extract content from specifically docx format
    '''
    doc = Document(io.BytesIO(body))
    texts = []

    # paragraphs + manual line breaks
    for p in doc.paragraphs:
        line = ""
        for run in p.runs:
            line += run.text
            # add manual line breaks
            for br in run.element.findall('.//w:br', run.element.nsmap):
                line += "\n"
        texts.append(line)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)

    return "\n".join(texts)

def get_text_from_s3(object: str):
    """
    Fetch text from S3 for many file formats:
    PDF, DOCX, DOC, TXT, CSV, JSON
    
    Parameter: 
    object: filename (make sure the file is first level to be accessed at bucket_name/filename)
    """

    bucket_object = f'legaldocstorage/{object}'
    bucket_name, object_name = bucket_object_separator(bucket_object)

    try:
        response = s3.get_object(Bucket=bucket_name, Key=object_name)
        body = response["Body"].read()
    except NoCredentialsError:
        raise Exception("❌ AWS credentials not available")

    ext = object_name.split(".")[-1].lower()

    # ------------------ PDF ------------------
    if ext == "pdf":
        try:
            text = ""
            with pdfplumber.open(io.BytesIO(body)) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
            return text
        except:
            # fallback
            reader = PyPDF2.PdfReader(io.BytesIO(body))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

    # ------------------ DOCX ------------------
    elif ext == "docx":
        return extract_docx_all_text(body)

    # ------------------ DOC ------------------
    elif ext == "doc":
        try:
            result = mammoth.extract_raw_text(io.BytesIO(body))
            return result.value
        except:
            # fallback: convert with docx2txt
            with tempfile.NamedTemporaryFile(suffix=".doc") as temp_doc:
                temp_doc.write(body)
                temp_doc.flush()
                return docx2txt.process(temp_doc.name)

    # ------------------ JSON ------------------
    elif ext == "json":
        return json.dumps(json.loads(body.decode("utf-8")), ensure_ascii=False, indent=2)

    # ------------------ CSV ------------------
    elif ext == "csv":
        return body.decode("utf-8")

    # ------------------ TXT-like ------------------
    elif ext in ["txt", "log", "md"]:
        return body.decode("utf-8", errors="ignore")


    # ------------------ XML ------------------
    elif ext == "xml":
        return body.decode("utf-8", errors="ignore")

    # # ------------------ RTF ------------------
    # elif ext == "rtf":
    #     import striprtf
    #     return striprtf.rtf_to_text(body.decode("utf-8", errors="ignore"))

    # # ------------------ ODT ------------------
    # elif ext == "odt":
    #     import pypandoc
    #     return pypandoc.convert_text(body.decode("utf-8", errors="ignore"), "plain", format="odt")

    # # ------------------ EPUB ------------------
    # elif ext == "epub":
    #     import ebooklib
    #     from ebooklib import epub
    #     from bs4 import BeautifulSoup
    #     book = epub.read_epub(io.BytesIO(body))
    #     text = []
    #     for item in book.get_items():
    #         if item.get_type() == ebooklib.ITEM_DOCUMENT:
    #             soup = BeautifulSoup(item.get_body_content(), "html.parser")
    #             text.append(soup.get_text())
    #     return "\n".join(text)

    # # ------------------ Image OCR ------------------
    # elif ext in ["png", "jpg", "jpeg", "tiff"]:
    #     import pytesseract
    #     from PIL import Image
    #     img = Image.open(io.BytesIO(body))
    #     return pytesseract.image_to_string(img, lang="eng")

    # else:
    #     raise Exception(f"❌ Unsupported file type: {ext}")

def download_s3_to_temp(bucket_object: str, suffix: str = None) -> str:
    """
    Download an S3 object to a temporary file and return the local path.

    Args:
        bucket_object (str): "bucket/key" format (e.g., "legaldocstorage/sample.pdf")
        suffix (str): Optional suffix (e.g., ".pdf"). If not provided, 
                      it will be inferred from the object name.

    Returns:
        str: Path to the temporary local file
    """
    bucket_name, object_name = bucket_object_separator(bucket_object)

    # infer suffix if not provided
    if suffix is None:
        ext = os.path.splitext(object_name)[-1]
        suffix = ext if ext else ".tmp"

    # create temp file
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp_file.close()  # close so boto3 can write into it

    # download from S3
    download_file_from_s3(bucket_object, tmp_file.name)

    return tmp_file.name

#For AWS MySQL RDS
def dml_mysql(query:str):
    '''
    Function to interact with MySQL db hosted by AWS, used for DDL and DML only
    '''
    conn = pymysql.connect(
    host='',
    user='admin',
    password='',
    database='',
    port=3306,
    # ssl={'ca': r'D:/Study/Education/Projects/Group_Project/secrets/us-east-1-bundle.pem'}
    )
    
    cursor = conn.cursor()
    
    cursor.execute(query)
    
    conn.close()

def query_mysql(query:str):
    '''
    Function to query the MySQL db, used for Query and retrieving data
    '''
    conn = pymysql.connect(
    host='',
    user='admin',
    password='',
    database='',
    port=3306,
    # ssl={'ca': r'D:/Study/Education/Projects/Group_Project/secrets/us-east-1-bundle.pem'}
    )
    
    cursor = conn.cursor()

    cursor.execute(query)
    result = cursor.fetchall()

    conn.close()
    
    return result

#For neo4j
URI = os.getenv('NEO4J_URI')
AUTH = ("neo4j", os.getenv('NEO4J_AUTH'))

def query_neo4j(query: str, **params):
    """
    Function to query the Neo4j db, used for Query and retrieving data
    """
    with GraphDatabase.driver(URI, auth=AUTH) as driver:
        driver.verify_connectivity()
        records, summary, keys = driver.execute_query(
            query,
            **params,  #Expect params to be in dictionary format
            database="neo4j"
        )
        return [record.data() for record in records]

def dml_ddl_neo4j(query: str, **params):
    """
    Function for DDL and DML in Neo4j database
    """    
    with GraphDatabase.driver(URI, auth=AUTH) as driver:
        driver.verify_connectivity()
        records, summary, keys = driver.execute_query(
            query,
            **params,   #Expect params to be in dictionary format
            database="neo4j"
        )

        print("Created {nodes_created} nodes, {rels_created} rels in {time} ms.".format(
            nodes_created=summary.counters.nodes_created,
            rels_created=summary.counters.relationships_created,
            time=summary.result_available_after
        ))

#File format changes
def save_to_txt(text: str, file_name: str):
    '''
    Save text string to a txt file
    
    Parameter:
    text: str
    file_name: local path to target file
    '''
    
    with open(f"D:/Study/Education/Projects/Group_Project/source/document/text_format/{file_name}.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("File saved as output.txt")
    
def docx_to_pdf(input_path, output_path=None):
    """
    Convert a .docx file to PDF automatically.
    
    Args:
        input_path (str): Path to the .docx file.
        output_path (str, optional): Path to save the converted PDF.
                                    Defaults to same folder as input.
                                    
    Warning: Auto deletes the source docx file after conversion, cannot be undone
    """
    if not input_path.lower().endswith(".docx"):
        raise ValueError("Only .docx files are supported by docx2pdf")

    # If no output path is specified, use the same directory
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + ".pdf"

    convert(input_path, output_path)
    os.remove(input_path)
    
    print(f"✅ Converted: {input_path} → {output_path} and safely deleted {input_path}")

def doc_to_pdf(input_dir = 'D:/Study/Education/Projects/Group_Project/source/document/original_doc', output_dir='D:/Study/Education/Projects/Group_Project/source/document'):
    """
    Convert all .doc files to PDF from a whole folder to a new folder
    
    Warning: Auto deletes the source docx file after conversion, cannot be undone
    """
    libreoffice_path = r"C:/Program Files/LibreOffice/program/soffice.exe"

    if not os.path.exists(libreoffice_path):
        raise FileNotFoundError("LibreOffice not found. Please install it or update the path.")

    input_files = os.listdir(input_dir)

    if len(input_files)>0:
        for input_path in input_files:
            input_full_path = f'{input_dir}/{input_path}'
            subprocess.run([
                libreoffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                input_full_path
            ], check=True)

            pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_full_path))[0] + ".pdf")
            os.remove(input_full_path)
            print(f"✅ Converted: {input_full_path} → {pdf_path} and safely deleted {input_full_path}")
    return pdf_path

def doc_to_docx(input_path, output_path=None):
    '''
    Convert doc to docx file for easier parsing
    
    Parameter:
    
    input_path: path to doc file
    
    output_path: path to docx target file
    '''
    # Convert to absolute normalized Windows path
    input_path = os.path.abspath(input_path)
    input_path = input_path.replace("/", "\\")

    if output_path is None:
        base, _ = os.path.splitext(input_path)
        output_path = base + ".docx"

    output_path = os.path.abspath(output_path)
    output_path = output_path.replace("/", "\\")

    # Ensure output folder exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Convert using Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False

    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=16)  # 16 = wdFormatXMLDocument
    doc.Close()
    word.Quit()

    return output_path
